import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
import subprocess

def sanitize_text(text):
    # Replace non-XML-compatible characters with a placeholder
    return re.sub(r'[^\x09\x0A\x0D\x20-\x7E]', '?', text)

def add_header(document):
    # Add Header
    header_section = document.sections[0]
    header = header_section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the logo
    logo_path = "Logo Horizontal.png"
    logo_height = Cm(1.5)
    run = header_paragraph.add_run()
    run.add_picture(logo_path, height=logo_height)

    # Add space after header
    header.add_paragraph()

def create_cover_page(document, keyword):
    add_header(document)
    # Add the logo image
    document.add_picture('tri.png', width=Inches(2))
    document.paragraphs[-1].alignment = 1  # Center align the image
   
    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    while len(document.paragraphs) % 4 != 0:
        document.add_paragraph()

    title = """Sample Penetration Test Report Example Company"""
    document.add_heading(title, level=0).alignment = 1 # Center align the text

    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    while len(document.paragraphs) % 14 != 0:
        document.add_paragraph()

    # Add details
    text = f"""
    Company: {keyword}
    Date: {date.today().strftime('%d %B %Y')}
    Version 1.0"""
    document.add_paragraph(text)

def generate_report(keyword, scan_result, exploit_result, output_dir):
    document = Document()
    create_cover_page(document, keyword)
    # Introduction
    document.add_heading('Introduction', level=1)
    introduction = 'This is a pentesting report for CVE-2018-14847. ' \
                   'CVE-2018-14847 is a vulnerability in the Winbox component ' \
                   'of MikroTik RouterOS. The vulnerability allows an unauthenticated ' \
                   'remote attacker to execute arbitrary code on vulnerable routers. ' \
                   'The impact of this vulnerability includes unauthorized access ' \
                   'to the router, data theft, and potential network compromise. ' \
                   'It is crucial to address this vulnerability promptly to protect ' \
                   'the security and integrity of the network infrastructure.'
    paragraph = document.add_paragraph(introduction)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Executive Summary
    vulnerable_devices = 0
    exploitable_devices = 0
    extracted_credentials = 0
    ipv4_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b' 
    with open(scan_result, 'r') as f:
        for line in f:
            if re.search(ipv4_pattern, line):
                vulnerable_devices += 1
    with open(exploit_result, 'r') as f:
        for line in f:
            if 'Exploit successful' in line:
                exploitable_devices += 1
    with open(exploit_result, 'r') as f:
        for line in f:
            if 'User' in line:
                extracted_credentials += 1
    document.add_heading('Executive Summary', level=1)
    document.add_paragraph(f'Number of Vulnerable Devices: {vulnerable_devices}')
    document.add_paragraph(f'Number of Exploitable Devices: {exploitable_devices}')
    document.add_paragraph(f'Number of Extracted Credentials: {extracted_credentials}')
    executivesummary_explanation = 'The executive summary provides an overview of the findings ' \
                                    'and key statistics related to the penetration testing performed. ' \
                                    'It highlights the number of vulnerable devices, exploitable devices, ' \
                                    'and the count of extracted credentials. These metrics help in ' \
                                    'assessing the severity and impact of the identified vulnerability ' \
                                    'and provide an initial understanding of the security posture of ' \
                                    'the network infrastructure.'
    paragraph = document.add_paragraph(executivesummary_explanation)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    # Findings
    document.add_heading('Findings', level=1)
    
    if vulnerable_devices == 0:
        document.add_heading('No Vulnerable Devices', level=2)
        document.add_paragraph('No vulnerable devices and exploitable devices were found. Your network infrastructure is not vulnerable to CVE-2018-14847.')
    else:
        # Vulnerable Devices
        document.add_heading('Vulnerable Devices', level=2)
        document.add_paragraph(f'Detected devices that are vulnerable to CVE-2018-14847:')
        with open(scan_result, 'r', encoding='utf-8', errors='ignore') as f:
            scan_result_content = f.read()
            sanitized_content = sanitize_text(scan_result_content)
            paragraph = document.add_paragraph(sanitized_content)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in paragraph.runs:
                run.font.size = Pt(8)

        # Exploitable Devices
        document.add_heading('Exploitable Devices', level=2)
        document.add_paragraph(f'Credential dumped from exploitable devices:')
        with open(exploit_result, 'r', encoding='utf-8', errors='ignore') as f:
            exploit_result_content = f.read()
            sanitized_content = sanitize_text(exploit_result_content)
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.text = sanitized_content
            run.font.size = Pt(10)

    # Recommendation
    document.add_heading('Recommendation', level=1)
    if vulnerable_devices == 0:
        recommendation = 'It is still recommended to keep the security posture monitored, maintained, and improved to ensure the ongoing protection ' \
                         'and integrity of your network infrastructure.'
                         
    else:
        recommendation = 'To mitigate CVE-2018-14847, it is recommended to take the following steps:\n' \
                         '- Update the MikroTik RouterOS to the latest version available.\n' \
                         '- Apply firewall rules to restrict access to the Winbox service.\n' \
                         '- Regularly monitor and review logs for any suspicious activity.\n' \
                         '- Implement strong password policies and avoid using default credentials.\n' \
                         '- Keep the network infrastructure and devices up to date with security patches.\n' \
                         '- Conduct regular security assessments and penetration testing to identify vulnerabilities.'
    document.add_paragraph(recommendation)

    # Save the report as a DOCX
    docx_report_path = os.path.join(output_dir, f"{os.path.basename(output_dir)}_report.docx")
    document.save(docx_report_path)

    # Convert the DOCX report to PDF using unoconv
    pdf_report_path = os.path.join(output_dir, f"{os.path.basename(output_dir)}_report.pdf")
    subprocess.run(["unoconv", "-f", "pdf", "-o", pdf_report_path, docx_report_path], check=True, stderr=subprocess.DEVNULL)

    # Remove the temporary DOCX report file
    os.remove(docx_report_path)

    return pdf_report_path
